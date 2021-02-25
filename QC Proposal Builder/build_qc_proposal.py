'''
QC Proposal Builder
Author: Ryan Carpus
Please close this file and the the read_me.txt file in the READ_ME folder!
Do not modify this file unless you know what you are doing.
Modify values in the config.py file
'''

import config as cf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
import datetime
from docx.oxml.shape import CT_Inline
from docx.shared import Inches, Cm, Pt
from openpyxl import Workbook, load_workbook

#Credit for this function definition to zhongle at https://github.com/python-openxml/python-docx/issues/235
def add_picture_to_run(run, picture_file, width=None, height=None):
    """
    Add a picture at the end of a run.
    """
    rId, image = run.part.get_or_add_image(picture_file)
    cx, cy = image.scaled_dimensions(width, height)
    shape_id, filename = run.part.next_id, image.filename
    inline = CT_Inline.new_pic_inline(shape_id, rId, filename, cx, cy)
    run._r.add_drawing(inline)

#opens a blank document
document = Document('qc_proposal_template.docx')

#Gets the current date and formats it. See link for how this works
#https://learnandlearn.com/python-programming/python-reference/python-get-current-date
d = datetime.datetime.today().strftime('%B %d, %Y')

#Header
document.sections[0].header.paragraphs[1].text = 'G2 Project No. ' + cf.project_number

#Set margins
sections = document.sections
for section in sections:
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.bottom_margin = Inches(1.0)
    section.top_margin = Inches(0.7)
    section.gutter = 0

#Set font
font = document.styles['Normal'].font
font.name = 'Lucida Sans'
font.size = Pt(10)

#buffer space because I am an idiot and I don't understand styles)
for i in range(7):
    document.add_paragraph()

#date
document.add_paragraph(d)
document.add_paragraph()

#Client information
if cf.client_suffix != None:
    document.add_paragraph(cf.client_prefix + ' ' + cf.client_first_name + ' ' + cf.client_last_name + ', ' + cf.client_suffix)
else:
    document.add_paragraph(cf.client_prefix + ' ' + cf.client_first_name + ' ' + cf.client_last_name)
document.add_paragraph(cf.client_company)
document.add_paragraph(cf.client_street_address)
document.add_paragraph(cf.client_city + ', ' + cf.client_state + ' ' + cf.client_zip)
document.add_paragraph()

#Job information
document.add_paragraph('RE:\tProposal for Construction Engineering Services').paragraph_format.tab_stops.add_tab_stop(Inches(.5))
document.add_paragraph('\t' + cf.project_name).paragraph_format.tab_stops.add_tab_stop(Inches(.5))
document.add_paragraph('\t' + cf.project_street_address).paragraph_format.tab_stops.add_tab_stop(Inches(.5))
document.add_paragraph('\t' + cf.project_city + ', ' + cf.project_state + ' ' + cf.project_zip).paragraph_format.tab_stops.add_tab_stop(Inches(.5))
document.add_paragraph()

#Dear client
document.add_paragraph('Dear ' + cf.client_prefix + ' ' + cf.client_last_name + ':')
document.add_paragraph()

#Introduction
intro = document.add_paragraph()
intro.add_run('In accordance with your request, we present our Proposal for Construction Engineering Services for the '
              + cf.project_name + ' project in ' + cf.project_city + ', ' + cf.project_state + '.  '
              + 'We understand the project will consist of ')
intro.add_run(cf.project_description)
document.add_paragraph()

#Convert list of inspection services to a readable string
if len(cf.work_types) == 1:
    service = cf.work_types[0].name + ' inspection.'
elif len(cf.work_types) == 2:
    service = cf.work_types[0].name + ' and ' + cf.work_types[1].name + ' inspection.'
elif len(cf.work_types) > 2:
    service = ''
    for i in range(len(cf.work_types)):
        if len(cf.work_types) - (i + 1) == 0:
            service = service + ', and ' + cf.work_types[i].name + ' inspection.'
        elif i == 0:
            service = cf.work_types[i].name
        else:
            service = service + ', ' + cf.work_types[i].name
#Write Special inspections paragraph
special_inspections_statement = document.add_paragraph()
special_inspections_statement.add_run('Based on our review of the drawings, we anticipate special inspection services will include '
                                      + service)

#Write personnel paragraph and scope of services intro
document.add_paragraph(cf.personnel_header, style='Heading 1')
document.add_paragraph(cf.personnel_body)
document.add_paragraph(cf.scope_header, style='Heading 1')
document.add_paragraph(cf.scope_body)

#Write out bullet points for each item in scope
for each in cf.work_types:
    document.add_paragraph(each.heading, style='heading 1')
    for line in each.body:
        document.add_paragraph(line, style='List Bullet')

#Professional fees
document.add_paragraph('PROFESSIONAL FEES', style='Heading 1')
fees = document.add_paragraph()
fees.add_run('Our staff has extensive experience providing construction observation and quality control testing services for similar commercial projects.  '
             + 'In general, we charge fees for our field services as outlined above based on a ')
fees.add_run('daily rate of $' + cf.full_day + '.00 (up to 8 hours portal to portal) or half day rate of $'
             + cf.half_day + '.00 (up to 4 hours portal to portal)').bold=True
fees.add_run('.  Overtime hours (in excess of 8 hours) will be charged at a rate of $' + cf.overtime_hour
             + ' per hour.  Weekends and holidays will be charged at a 50 percent premium.')
document.add_paragraph('Our daily rate includes the costs for the project as follows:')
document.add_paragraph('On site observation and testing services by our field engineer', style = 'List Bullet')
document.add_paragraph('Project management time (excluding special meetings)', style = 'List Bullet')
document.add_paragraph('Administrative fees for processing written observation and test reports', style = 'List Bullet')
document.add_paragraph('Engineering review of test reports', style = 'List Bullet')
document.add_paragraph('Equipment and material charges', style = 'List Bullet')

document.add_paragraph('We charge for laboratory tests according to the rate schedule provided below:')
document.add_paragraph()

#loads lab rate sheet
wb = load_workbook('lab_unit_rates.xlsx', data_only=True)
ws = wb.active
lab_rate_data = []
for row in ws.values:
    lab_rate_data.append(row)

lab_rate_table = document.add_table(rows=0, cols=2, style='GridTable1Light-Accent1')
for i in range(len(lab_rate_data)):
    cells = lab_rate_table.add_row().cells
    cells[0].text = str(lab_rate_data[i][0])
    if i == 0:
        cells[1].text = str(lab_rate_data[i][1])
    else:
        cells[1].text = '$' + str(lab_rate_data[i][1]) + '.00'
    
#formats table column width
widths = [Inches(6.25),Inches(1.25)]
for row in lab_rate_table.rows:
    for i, w in enumerate(widths):
        row.cells[i].width = w
        
#formats alignment
for i in range(1,len(lab_rate_table.rows)):
    lab_rate_table.rows[i].cells[0].paragraphs[0].runs[0].bold=False
    lab_rate_table.rows[i].cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
document.add_paragraph()

document.add_paragraph('Based on our review of the drawings, following is an estimate of our fees for construction observation and testing services:')
document.add_paragraph()

#Loads cost estimate spreadsheet and pastes into document
wb = load_workbook('estimate_sheet.xlsx', data_only=True)
ws = wb.active
estimate_data = []
for row in ws.values:
    estimate_data.append(row)

estimate_table = document.add_table(rows=0, cols=5, style='GridTable1Light-Accent1')
for i in range(len(estimate_data)):
    cells = estimate_table.add_row().cells
    for j in [0,1,2]:
        if str(estimate_data[i][j]) != 'None':
            cells[j].text = str(estimate_data[i][j])
    for j in [3,4]:
        if i > 0 and str(estimate_data[i][j]) != 'None':
            cells[j].text = '$' + str(estimate_data[i][j]) + '.00'
        elif str(estimate_data[i][j]) != 'None':
            cells[j].text = str(estimate_data[i][j])
cells[4].paragraphs[0].runs[0].bold=True

#formats table column width
widths = [Inches(2),Inches(2),Inches(1),Inches(1),Inches(1)]
for row in estimate_table.rows:
    for i, w in enumerate(widths):
        row.cells[i].width = w
#formats alignment
for i in range(1,len(estimate_table.rows)):
    for j in range(2,5):
        estimate_table.rows[i].cells[j].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

document.add_paragraph()
document.add_paragraph(cf.budget_explanation)
document.add_paragraph(cf.terms_and_conditions_header, style='Heading 2')
document.add_paragraph(cf.terms_and_conditions_body)
document.add_paragraph()
document.add_paragraph('Sincerely,')
document.add_paragraph()
G2 = document.add_paragraph()
G2.add_run('Our Company Name Which Is Hard Coded For Some Reason').bold = True

#Writes engineer signatures
if cf.engineer_1_signature != None:
    signature_line = document.add_paragraph().add_run()
    add_picture_to_run(signature_line, cf.engineer_1_signature, height=Inches(0.75))
    if cf.engineer_2_signature != None:
        add_picture_to_run(signature_line, cf.engineer_2_signature, height=Inches(0.75))

#Writes engineer names
names_line = document.add_paragraph()
names_line.paragraph_format.tab_stops.add_tab_stop(Inches(3.5))
names_line.add_run(cf.engineer_1_name + '\t' + cf.engineer_2_name)
#Write engineer titles
titles_line = document.add_paragraph()
titles_line.paragraph_format.tab_stops.add_tab_stop(Inches(3.5))
titles_line.add_run(cf.engineer_1_title + '\t' + cf.engineer_2_title)

document.add_paragraph(cf.initials)
document.add_paragraph()
document.add_paragraph('Encl:\tGeneral Conditions').paragraph_format.tab_stops.add_tab_stop(Inches(0.5))
document.add_paragraph('\tFee Schedule').paragraph_format.tab_stops.add_tab_stop(Inches(0.5))
document.add_paragraph()
accepted = document.add_paragraph()
accepted.add_run(('Accepted for ' + cf.client_company + ':').upper()).bold = True
document.add_paragraph('BY:\t________________________________________')
document.add_paragraph('DATE:\t________________________________________')


document.save(cf.document_name)



